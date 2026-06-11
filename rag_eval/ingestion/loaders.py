from __future__ import annotations

import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Dict, Optional


SUPPORTED_EXTENSIONS = {".txt", ".md", ".html", ".htm", ".pdf", ".docx"}
URL_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/125.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}
PLAYWRIGHT_INSTALL_HINT = (
    "浏览器渲染依赖未就绪；请先执行 "
    ".venv/bin/python -m playwright install chromium"
)
PLAYWRIGHT_DEPENDENCY_HINT = (
    "浏览器渲染依赖未安装；请先运行 pip install -r requirements.txt，"
    "然后执行 .venv/bin/python -m playwright install chromium"
)


@dataclass
class ParsedDocument:
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self.title = ""
        self._in_title = False
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag == "title":
            self._in_title = True
        if tag in {"p", "br", "div", "section", "article", "li", "h1", "h2", "h3"}:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"} and self._skip_depth > 0:
            self._skip_depth -= 1
        if tag == "title":
            self._in_title = False
        if tag in {"p", "div", "section", "article", "li"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._in_title:
            self.title += data.strip()
        self._parts.append(data)

    def text(self) -> str:
        return normalize_text(" ".join(self._parts))


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _parse_html(content: str, metadata: Dict[str, Any]) -> ParsedDocument:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "html.parser")
        title = soup.title.get_text(" ", strip=True) if soup.title else ""
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = normalize_text(soup.get_text("\n"))
    except Exception:
        extractor = _TextExtractor()
        extractor.feed(content)
        title = extractor.title
        text = extractor.text()

    meta = dict(metadata)
    if title:
        meta["title"] = title
    return ParsedDocument(content=text, metadata=meta)


def _parse_pdf(path: Path, metadata: Dict[str, Any]) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF parsing requires pypdf. Install requirements.txt first.") from exc

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    meta = dict(metadata)
    meta["page_count"] = len(reader.pages)
    return ParsedDocument(content=normalize_text("\n\n".join(pages)), metadata=meta)


def _parse_docx(path: Path, metadata: Dict[str, Any]) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("DOCX parsing requires python-docx. Install requirements.txt first.") from exc

    document = DocxDocument(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ParsedDocument(content=normalize_text("\n".join(parts)), metadata=dict(metadata))


def _load_playwright():
    try:
        from playwright.sync_api import Error as PlaywrightError
        from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise RuntimeError(PLAYWRIGHT_DEPENDENCY_HINT) from exc
    return sync_playwright, PlaywrightError, PlaywrightTimeoutError


def _render_url_text(url: str, timeout: int = 20) -> tuple[str, str]:
    sync_playwright, playwright_error, playwright_timeout_error = _load_playwright()
    timeout_ms = timeout * 1000
    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(headless=True)
            try:
                page = browser.new_page(
                    user_agent=URL_HEADERS["User-Agent"],
                    locale="zh-CN",
                    extra_http_headers={
                        "Accept-Language": URL_HEADERS["Accept-Language"],
                    },
                )
                try:
                    page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
                except playwright_timeout_error:
                    pass
                try:
                    page.wait_for_load_state("networkidle", timeout=min(timeout_ms, 5000))
                except playwright_timeout_error:
                    pass
                title = page.title()
                text = page.evaluate("() => document.body ? document.body.innerText : ''")
                return title or "", normalize_text(text or "")
            finally:
                browser.close()
    except RuntimeError:
        raise
    except playwright_error as exc:
        message = str(exc)
        if "Executable doesn't exist" in message or "playwright install" in message:
            raise RuntimeError(PLAYWRIGHT_INSTALL_HINT) from exc
        raise RuntimeError(f"浏览器渲染 URL 失败：{message}") from exc


def parse_file(path: str | Path, original_name: Optional[str] = None) -> ParsedDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    name = original_name or file_path.name
    metadata: Dict[str, Any] = {
        "source": name,
        "path": str(file_path),
        "extension": suffix,
        "title": file_path.stem,
    }

    if suffix == ".doc":
        raise ValueError("legacy .doc is not supported; please convert it to .docx")
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported file type: {suffix or '<none>'}")
    if not file_path.exists():
        raise FileNotFoundError(str(file_path))

    if suffix in {".txt", ".md"}:
        return ParsedDocument(content=normalize_text(_read_text_file(file_path)), metadata=metadata)
    if suffix in {".html", ".htm"}:
        return _parse_html(_read_text_file(file_path), metadata)
    if suffix == ".pdf":
        return _parse_pdf(file_path, metadata)
    if suffix == ".docx":
        return _parse_docx(file_path, metadata)
    raise ValueError(f"unsupported file type: {suffix}")


def parse_url(url: str, timeout: int = 20) -> ParsedDocument:
    title, text = _render_url_text(url, timeout=timeout)
    metadata = {"source": url, "url": url, "extension": ".html", "title": title or url}
    return ParsedDocument(content=text, metadata=metadata)
